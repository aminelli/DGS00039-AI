""" per python 2.x """
""" import docx

def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print (readtxt('Python\\Text Processing\\asset\\Tutorialspoint.docx')) """

""" per python 3.x """
from docx import Document

def readtxt(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print (readtxt('Python\\Text Processing\\asset\\Tutorialspoint.docx'))
