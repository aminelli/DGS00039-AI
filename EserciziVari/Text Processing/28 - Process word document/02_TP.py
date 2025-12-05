""" per python 2.x """
""" import docx

doc = docx.Document('Python\\Text Processing\\asset\\Tutorialspoint.docx')
print (len(doc.paragraphs))

print (doc.paragraphs[2].text) """

""" per python 3.x """
from docx import Document

doc = Document('Python\\Text Processing\\asset\\Tutorialspoint.docx')
print (len(doc.paragraphs))

print (doc.paragraphs[2].text)